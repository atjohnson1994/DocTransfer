# Import dependencies
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from zipfile import ZipFile
from docx.oxml import OxmlElement
from docx.oxml import parse_xml

# =============================== SECTION 1: Style mapping =============================================
# Purpose: Defines the "style" for each paragraph based on the "style" in the original formatting.

style_mapping = {
    "Heading 1": "Heading 1",
    "Heading 2": "Heading 2",
    'Heading 3': 'Heading 3',
    'Heading 4': 'Heading 4',
    'Heading 5': 'Heading 5',
    'Heading 6': 'Heading 6',
    "Normal": "00_TEXT",
    "List Paragraph": "00_BULLET",
    "Documet_Title": "00_TEXT",
    'Norm_NoIndent': "00_TEXT",
    '_ArialTableHeader': '00_TITLE TABLE',
    '00_TITLE TABLE': '00_TITLE TABLE',
    '_TimesTableBody': '00_TEXT',
    'Caption': '00_PICTURE',
    'Norm_1_2Head': '00_TEXT',
    'Norm_4Head': '00_TEXT',
    'Norm_3Head': '00_TEXT',
    '00_TEXT': '00_TEXT'
}


# =============================== SECTION 2: Content Extraction =============================================
# Purpose: Extract text, paragraph properties, images, tables, and header/footer from source document.
# Functions: 2.1 : is_paragraph_in_list
#            2.2 : extract_revision_text
#            2.3 : extract_approval_text
#            2.4 : extract_document_information
#            2.5 : extract_content_with_details
#            2.6 : extract_and_copy_tables
#            2.7 : extract_images_from_docx

# (2.1) Determine if a paragraph is part of a list
def is_paragraph_in_list(paragraph):
    """Check if a paragraph is part of a bullet or numbered list by inspecting its XML properties."""
    num_pr = paragraph._p.xpath('.//w:numPr')
    return bool(num_pr)



# (2.2) Extract revision text
def extract_revision_text(source_doc_path, first_cell_text="Revision History"):

    source_doc = Document(source_doc_path)

    for table in source_doc.tables:
        if table.cell(0, 0).text.strip() == first_cell_text:
            return [[cell.text.strip() for cell in row.cells] for row in table.rows]

    return None  # Return None if no matching table is found



# (2.3) Extract approval text
def extract_approval_text(source_doc_path, first_cell_text="Approval Table"):

    source_doc = Document(source_doc_path)

    for table in source_doc.tables:
        if table.cell(0, 0).text.strip() == first_cell_text:
            return [[cell.text.strip() for cell in row.cells] for row in table.rows]

    return None  # Return None if no matching table is found


# (2.4)
def extract_document_information(source_doc_path):
    doc = Document(source_doc_path)

    header_content = None  
    footer_content = None  

    # Locate header and footer
    for section in doc.sections:
        header = section.header
        footer = section.footer

        # Extract header
        try:
            for table in header.tables:
                header_content = [[cell.text.strip() for cell in row.cells] for row in table.rows]
                break  # Exit loop if header content found
        except:
            pass  # Skip header extraction if it fails

        # Extract footer
        try:
            for table in footer.tables:
                footer_content = [[cell.text.strip() for cell in row.cells] for row in table.rows]
                break  # Exit loop if footer content found
        except:
            pass  # Skip footer extraction if it fails

    return [header_content, footer_content]


# (2.5)
def extract_content_with_details(source_doc_path):
    source_doc = Document(source_doc_path)
    content = []
    paragraph_idx = 0  # Index to track paragraphs
    table_idx = 0      # Index to track tables

    for block in source_doc.element.body:
        if block.tag.endswith("p"):  # Paragraph
            if paragraph_idx < len(source_doc.paragraphs):
                p_obj = source_doc.paragraphs[paragraph_idx]
                is_list = is_paragraph_in_list(p_obj)
                content.append({
                    "type": "paragraph",
                    "text": p_obj.text,
                    "style": p_obj.style.name,
                    "is_list": is_list,
                    "runs": [
                        {
                            "text": run.text,
                            "bold": run.bold,
                            "italic": run.italic
                        }
                        for run in p_obj.runs
                    ]
                })
                paragraph_idx += 1
        elif block.tag.endswith("tbl"):  # Table
            if table_idx < len(source_doc.tables):
                table_obj = source_doc.tables[table_idx]
                table_data = []
                for row in table_obj.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        cell_style = cell.paragraphs[0].style.name if cell.paragraphs else None
                        row_data.append({
                            "text": cell_text,
                            "style": cell_style
                        })
                    table_data.append(row_data)
                content.append({
                    "type": "table",
                    "data": table_data
                })
                table_idx += 1

    return content


# (2.6)
def extract_and_copy_tables(source_doc_path, output_folder):
    # Open the source document
    source_doc = Document(source_doc_path)


    # Create a new destination document
    dest_doc = Document()


    # Iterate over each table in the source document and copy it to the destination document
    for table in source_doc.tables:
        # Copy the table's XML and append it to the destination document
        table_xml = table._element.xml  # Extract full table XML
        new_table = parse_xml(table_xml)  # Parse into new table object
        dest_doc._element.body.append(new_table)  # Append to the document


    # Get the original document name and append "_supplemental_tables"
    base_name = os.path.splitext(os.path.basename(source_doc_path))[0][:6]
    destination_doc_name = f"{base_name}_supplemental_tables.docx"
    destination_doc_path = os.path.join(output_folder, destination_doc_name)


    # Save the destination document
    dest_doc.save(destination_doc_path)
    print(f"Tables extracted and saved to {destination_doc_path}")


# (2.7)
def extract_images_from_docx(docx_path, output_dir):
    """Extract images from a DOCX file and save them to a specified output directory."""
    os.makedirs(output_dir, exist_ok=True)


    # Open the DOCX file as a ZIP archive
    with ZipFile(docx_path, 'r') as docx_zip:
        # Extract all the relationships to identify the images
        for file in docx_zip.namelist():
            if file.startswith('word/media/'):
                # Extract the image file
                image_name = file.split('/')[-1]
                image_data = docx_zip.read(file)


                # Save the image to the output directory
                image_path = os.path.join(output_dir, image_name)
                with open(image_path, 'wb') as img_file:
                    img_file.write(image_data)
                print(f"Extracted image: {image_name}")


# =============================== SECTION 3: Document Creation =============================================
# Purpose: Transfer all extracted content into new template, apply formatting and styling, then save.
# Functions: 3.1 : input_document_information
#            3.2 : apply_paragraph_style
#            3.3 : center_cell_content
#            3.4 : set_font_size
#            3.5 : input_approvals_revisions_text
#            3.6 : write_content_with_existing_styles
#            3.7 : italicize_and_resize_caption_style
#            3.8 : insert_images_by_filename

# (3.1) Input title, doc number, revision
def input_document_information(finished_good, doc_information):
    try:
        doc = Document(finished_good)
    except Exception as e:
        print(f"Error opening document: {e}")
        return

    # Locate header and footer
    for section in doc.sections:
        header = section.header
        footer = section.footer

        try:
            for table in header.tables:
                try:
                    cell_0_5 = table.cell(0, 5)
                    cell_0_5.text = doc_information[0][0][5]
                    apply_paragraph_style(cell_0_5.paragraphs[0], "00_BOLD")
                    center_cell_content(cell_0_5)
                except (IndexError, AttributeError) as e:
                    print(f"Error setting header cell (0,5): {e}")
                
                try:
                    cell_1_4 = table.cell(1, 4)
                    cell_1_4.text = doc_information[0][1][4]
                    apply_paragraph_style(cell_1_4.paragraphs[0], "00_BOLD")
                    center_cell_content(cell_1_4)
                except (IndexError, AttributeError) as e:
                    print(f"Error setting header cell (1,4): {e}")
                
                try:
                    cell_2_3 = table.cell(2, 3)
                    cell_2_3.text = doc_information[0][2][3]
                    apply_paragraph_style(cell_2_3.paragraphs[0], "00_HEADER")
                    center_cell_content(cell_2_3)
                except (IndexError, AttributeError) as e:
                    print(f"Error setting header cell (2,3): {e}")
                
                try:
                    cell_3_0 = table.cell(3, 0)
                    cell_3_0.text = doc_information[0][3][0]
                    apply_paragraph_style(cell_3_0.paragraphs[0], "00_HEADER TITLE")
                    center_cell_content(cell_3_0)
                except (IndexError, AttributeError) as e:
                    print(f"Error setting header cell (3,0): {e}")
        except Exception as e:
            print(f"Error processing header tables: {e}")

        try:
            for table in footer.tables:
                try:
                    cell_0_2 = table.cell(0, 2)
                    cell_0_2.text = doc_information[1][0][2]
                    apply_paragraph_style(cell_0_2.paragraphs[0], "00_BOLD")
                    set_font_size(cell_0_2, 9)
                except (IndexError, AttributeError) as e:
                    print(f"Error setting footer cell (0,2): {e}")
        except Exception as e:
            print(f"Error processing footer tables: {e}")

    try:
        doc.save(finished_good)
    except Exception as e:
        print(f"Error saving document: {e}")


# (3.2)
def apply_paragraph_style(paragraph, style_name):
    try:
        # Try applying a paragraph style directly
        paragraph.style = style_name
    except ValueError:
        # If it's a character style, apply it to the run
        if len(paragraph.runs) > 0:
            run = paragraph.runs[0]  # Use the first run
            run.style = style_name  # Apply character style to the run
        else:
            # If no runs exist, add a run and apply the style
            run = paragraph.add_run(paragraph.text)
            run.style = style_name


# (3.3)
def center_cell_content(cell):
    # Center the content of each paragraph in the cell
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


# (3.4)
def set_font_size(cell, font_size):
    """Set the font size of all runs in the cell's paragraphs."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(font_size)



# (3.5) Function to input approvals and revision history
def input_approvals_revisions_text(finished_good, revision_history, approvals):
    doc = Document(finished_good)

    approval_count = 0
    revision_count = 0
    tables_to_remove = []


    for index, table in enumerate(doc.tables):
        if table.cell(0, 0).text.strip() == "Approval Table":
            approval_count += 1
            if approval_count == 2:  # Mark second occurrence for removal
                tables_to_remove.append(index)
            else:
                for i in range(4):
                    table.cell(2, i).text = approvals[2][i]
                    table.cell(2, i).paragraphs[0].style = "00_TEXT"


        if table.cell(0, 0).text.strip() == "Revision History":
            revision_count += 1
            if revision_count == 2:  # Mark second occurrence for removal
                tables_to_remove.append(index)
            else:
                for i in range(3):
                    table.cell(2, i).text = revision_history[2][i]
                    table.cell(2, i).paragraphs[0].style = "00_TEXT"


    # Remove tables marked for deletion
    for i in sorted(tables_to_remove, reverse=True):  
        tbl = doc.tables[i]._element
        tbl.getparent().remove(tbl)


    doc.save(finished_good)


# (3.6)
def write_content_with_existing_styles(content, destination_doc_path, finished_good):
    dest_doc = Document(destination_doc_path)

    # Check if the required styles exist in the destination document
    for item in content:
        if item["type"] == "paragraph":
            style_name = style_mapping.get(item["style"], None)
            if style_name:
                style_names = [s.name for s in dest_doc.styles]
                if style_name not in style_names:
                    raise ValueError(f"Style '{style_name}' not found in destination document.")

    # Write the content in the same order
    for item in content:
        if item["type"] == "paragraph":
            style_name = style_mapping.get(item["style"], "00_TEXT")  # Default to '00_TEXT' if not mapped
            paragraph = dest_doc.add_paragraph(style=style_name)

            # Only apply list style if explicitly marked as a list and not a heading
            if item.get("is_list", False) and "HEADING" not in style_name:
                paragraph.style = "00_BULLET"  # Use mapped bullet list style

            # Write runs with formatting
            for run_data in item["runs"]:
                run = paragraph.add_run(run_data["text"])
                run.bold = run_data["bold"]
                run.italic = run_data["italic"]

        elif item["type"] == "table":
            # Add a table to the destination document
            table_data = item["data"]
            if table_data:
                table = dest_doc.add_table(rows=0, cols=len(table_data[0]))
                table.style = 'Table Grid'  # Use a default table style

                # Populate table rows with styles
                for row_data in table_data:
                    row = table.add_row()
                    for idx, cell_data in enumerate(row_data):
                        cell = row.cells[idx]
                        cell.text = cell_data["text"]
                        # Map and apply styles
                        source_style = cell_data["style"]
                        dest_style = style_mapping.get(source_style, None)
                        if dest_style:
                            for paragraph in cell.paragraphs:
                                paragraph.style = dest_style
                        # Apply alignment to non-title cells
                        if dest_style != "00_TITLE TABLE":
                            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Save the modified document
    dest_doc.save(finished_good)
    print(f"Document saved to {finished_good}.")


# (3.7)
def italicize_and_resize_caption_style(finished_good):
    """Make all text with the 'Caption' style italicized and set font size to 9 in the DOCX document."""
    doc = Document(finished_good)
    for paragraph in doc.paragraphs:
        if paragraph.style.name == "00_PICTURE":
            for run in paragraph.runs:
                run.italic = True
                run.font.size = Pt(9)
        elif paragraph.text.startswith("Created from Template LLDC"):
            for run in paragraph.runs:
                run.italic = True
                run.font.size = Pt(8)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.save(finished_good)
    print(f"Text with 'Caption' style has been italicized and resized to 9 points in {finished_good}.")



# (3.8)
def insert_images_by_filename(destination_docx_path, image_folder):
    """Insert images into the DOCX file above their corresponding 'Figure X' text."""
    doc = Document(destination_docx_path)
    paragraphs = doc.paragraphs

    # Get a sorted list of images in the folder
    image_files = sorted(
        [f for f in os.listdir(image_folder) if f.lower().endswith(('png', 'jpg', 'jpeg'))]
    )

    for image_file in image_files:
        # Extract the figure number from the filename (e.g., "image2.png" -> "2")
        figure_number = ''.join(filter(str.isdigit, os.path.splitext(image_file)[0]))
        if not figure_number.isdigit():
            print(f"Skipping {image_file}: could not determine figure number.")
            continue

        figure_text = f"Figure {figure_number}"
        image_path = os.path.join(image_folder, image_file)

        # Search for the paragraph starting with "Figure X"
        for para in paragraphs:
            if para.text.startswith(figure_text):  # Checks if paragraph starts with "Figure X"
                # Insert image above the paragraph containing "Figure X"
                para_index = paragraphs.index(para)
                para_before = paragraphs[para_index].insert_paragraph_before()
                para_before.add_run().add_picture(image_path, width=Inches(3.0))  # Adjust width as needed
                para_before.alignment = WD_ALIGN_PARAGRAPH.CENTER
                break
        else:
            print(f"Figure {figure_number} not found in the document. Skipping {image_file}.")

    # Save the updated document
    doc.save(destination_docx_path)
    print("Images inserted successfully!")

    # Delete all images in the folder after completing
    for image_file in image_files:
        image_path = os.path.join(image_folder, image_file)
        os.remove(image_path)
        print(f"Deleted {image_file} from {image_folder}")




#======================================= SECTION 4: Function Calls ============================================
# Purpose: Call the functions from above to process all documents in the folder

# Directory paths
source_folder = 'Insert Non-Transferred Document Here'
destination_folder = 'resources/template.docx'
output_folder = 'Transferred Document Will Be Here'
image_folder = 'resources/extracted_images'

# Function to process each document in the source folder
def process_documents_in_folder():
    # Get all DOCX files in the source folder
    docx_files = [f for f in os.listdir(source_folder) if f.lower().endswith('.docx')]

    for docx_file in docx_files:
        source_doc_path = os.path.join(source_folder, docx_file)
        
        # Define the output path with the same name as the source file
        finished_good = os.path.join(output_folder, docx_file)

        # Extract Revision History
        revision_history = extract_revision_text(source_doc_path)

        # Save tables into supplemental document
        extract_and_copy_tables(source_doc_path, output_folder)
        
        # Extract Approvals Table
        approvals = extract_approval_text(source_doc_path)
        
        # Extract content from the source document
        content = extract_content_with_details(source_doc_path)
        
        # Extract document information
        doc_information = extract_document_information(source_doc_path)
        
        # Write content with styles to destination document
        write_content_with_existing_styles(content, destination_folder, finished_good)

        # Input revision and approvals
        input_approvals_revisions_text(finished_good, revision_history, approvals)

        # Input document information
        input_document_information(finished_good, doc_information)

        # Italicize and resize caption style
        italicize_and_resize_caption_style(finished_good)

        # Extract images from source DOCX
        extract_images_from_docx(source_doc_path, image_folder)

        # Insert images into the destination document
        insert_images_by_filename(finished_good, image_folder)

        print(f"Processed {docx_file} and saved to {finished_good}")

# Call the function to process the documents
process_documents_in_folder()
